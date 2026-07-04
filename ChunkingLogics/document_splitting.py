import os
import tempfile
from pathlib import Path

from langchain_text_splitters import MarkdownHeaderTextSplitter
from langchain_text_splitters import PythonCodeTextSplitter
from langchain_core.documents import Document 
from langchain_text_splitters import PythonCodeTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter, Language

import pymupdf
from docx import Document as Docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from unstructured.partition.pdf import partition_pdf
from unstructured.staging.base import elements_to_json


markdown_document = "# Foo\n\n    ## Bar\n\nHi this is Jim\n\nHi this is Joe\n\n ### Boo \n\n Hi this is Lance \n\n ## Baz\n\n Hi this is Molly"
markdown_table='''
| Product ID | Item Name | Category | Stock | Price | Status |
| :--- | :--- | :---: | ---: | ---: | :---: |
| P1001 | Wireless Mouse | Electronics | 45 | $29.99 | In Stock |
| P1002 | Mechanical Keyboard | Electronics | 12 | $89.50 | Low Stock |
| P1003 | Ergonomic Desk Chair | Furniture | 0 | $199.00 | Out of Stock |
| P1004 | USB-C Charging Cable | Accessories | 150 | $14.95 | In Stock |
'''

def read_code(path:str):
    with open(path, "r") as file:
        content = file.read()
        return(content)

def read_docx(path:str,single_text:bool):
    document = Docx(path)
    if single_text==False:
        return "\n".join([p.text for p in document.paragraphs])
    else:
        fulltext=""
        for paragraph in document.paragraphs:
            fulltext=fulltext+paragraph.text
        return fulltext
    
def read_pdf(path:str):
    """
    Load a pdf file and return it as a LangChain Document object.
    Args:
        file_path (str): The path to the text file to load.
        encoding (str): The encoding of the text file. Defaults to "utf-8".
    Returns:
        document object (Document): Page content and 
        metadata including source, file_name, file_extension, file_size, last_modified, created_time.
    Raises:
    """
    file_path = Path(path)
    doc = pymupdf.open(file_path)
    for page in doc:
        text = page.get_text()

    return [
        Document(
            page_content=text, 
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_extension": file_path.suffix,
                "file_size": file_path.stat().st_size, 
                "last_modified": file_path.stat().st_mtime, 
                "created_time": file_path.stat().st_ctime}
        )
    ]

def markdown_split():
    headers_to_split_on = [
    ("#", "Header 1"),
    ("##", "Header 2"),
    ("###", "Header 3"),
    ]

    table_to_slit=[
        ("|","column"),
        ("\n","new row"),
        (":---","column header ends")
    ]

    markdown_splitter = MarkdownHeaderTextSplitter(table_to_slit)
    md_header_splits = markdown_splitter.split_text(markdown_table)
    print(md_header_splits)

def python_split(code_data:str):
    code_splitter = PythonCodeTextSplitter(
    chunk_size=100, chunk_overlap=0,
    )

    return code_splitter.create_documents([code_data])

def any_code_split(code_data:str,language:Language):
    splitter=RecursiveCharacterTextSplitter.from_language(
        language=language,
        chunk_size=65,
        chunk_overlap=0,
    )

    docs=splitter.create_documents([code_data])
    return docs

def pdf_split(path:str):
    elements=partition_pdf(
        filename=path,

        strategy="hi_res",
        infer_table_structure=True,
        model_name="yolox"
    )
    return elements

if __name__ == "__main__":  
    #markdown_split()
    #file_path="test_files\Git Cheat Sheat.pdf"
    #docs=read_pdf(file_path)
    #print(docs)

    #data=read_docx("test_files\Effective_Threat_Modeling_using_TAM.docx",False)
    #print(data)

    #code=read_code("character_text_splitting.py")
    #print(code)
    #print("---------------------------------------------------")
    #for doc in python_split(code):
        #print(doc.page_content)

    #data=read_code("test_files\sample.json")
    #output=any_code_split(data,Language.JS)
    #for doc in output:
        #print(doc.page_content)

    output=pdf_split("test_files\Effective_Threat_Modeling_using_TAM.pdf")
    print(output)
    print(output[-4].metadata.text_as_html)